from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from pathlib import Path


OUTPUT = Path("/Users/mvc-m3/Documents/Marvin-Personal/CV/Marvin_Chang_CV_Eng_ABConvert_20260402.docx")


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    style.font.size = Pt(10.5)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Marvin Chang")
    run.bold = True
    run.font.size = Pt(20)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    p.add_run(
        "New Taipei City, Taiwan  |  maxio.mvc@gmail.com  |  +886-978-238-161\n"
        "Portfolio: mvc.softwidom.com  |  GitHub: github.com/softwise-mvc  |  LinkedIn: linkedin.com/in/mvc-profile/"
    )


def add_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Inches(0.25 * level)
    p.paragraph_format.space_after = Pt(0)
    p.add_run(text)


def add_role(doc: Document, title: str, company: str, dates: str, bullets: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{title} | {company}")
    r.bold = True
    p.add_run(f" | {dates}")
    for bullet in bullets:
        add_bullet(doc, bullet)


def build_cv() -> Document:
    doc = Document()
    set_doc_defaults(doc)
    add_title(doc)

    add_section_heading(doc, "Professional Summary")
    doc.add_paragraph(
        "AI-assisted system architect and builder with 20+ years across product, software architecture, and cross-functional delivery in software, fintech, blockchain, mobile, and IoT. "
        "I turn ambiguous ideas into working software systems, using Claude, Codex, and Gemini to accelerate design, debugging, workflow refinement, and iterative delivery. "
        "Over the past year, I have built and iterated AI-assisted production and lab systems with clear operating rules, verification layers, and fast feedback loops. Best suited for hands-on Technical Head / Engineering Lead roles where AI-assisted development, context quality, and shipping discipline matter."
    )

    add_section_heading(doc, "Core Competencies")
    add_bullet(doc, "AI-Assisted Software Delivery: multi-model workflow, context refinement, agent collaboration, debugging, and spec-to-code acceleration")
    add_bullet(doc, "System Architecture and Operational Guardrails: workflow design, runtime isolation, closed-loop review, QA gating, deployment verification")
    add_bullet(doc, "Hands-on Product Engineering: Next.js, TypeScript, Python, FastAPI, Supabase, PostgreSQL, Vercel, Docker")
    add_bullet(doc, "Technical Leadership: requirement shaping, release decisions, cross-functional execution, and delivery ownership")

    add_section_heading(doc, "Selected AI-Assisted Product Development / System Architecture")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Independent / Selected Portfolio Work")
    r.bold = True
    p.add_run(" | 2024 - Present")
    for bullet in [
        "Built multiple AI-assisted production and lab systems using Claude, Codex, and Gemini for architecture planning, implementation, bug triage, QA gating, deployment verification, and continuous improvement.",
        "Established multi-model workflow patterns with provider failover, audit logs, verification records, and runtime guardrails to reduce uncertainty in AI-assisted software delivery.",
        "Connected data collection, analysis, notification, simulation execution, and post-run review into closed-loop workflows running across Mac and Ubuntu environments.",
        "Collaborated on a confidential B2B digital banking mobile app project for iOS and Android, using AI-assisted development to accelerate requirement clarification, implementation planning, issue analysis, and delivery coordination.",
        "Worked closely with a colleague and AI tools to shorten feedback cycles across product logic, bug fixing, and technical decision-making.",
    ]:
        add_bullet(doc, bullet)

    add_section_heading(doc, "Selected Systems")
    add_bullet(doc, "TWST Main: Built a Taiwan stock analysis and signal platform combining market data, broker reports, and an LLM-based structured decision flow.")
    add_bullet(doc, "TWST Main: Added postmarket feedback loops, signal-quality controls, auditability, and OpenAI-to-Gemini failover to improve decision quality rather than only generate outputs.")
    add_bullet(doc, "TWST Lab: Built an isolated simulation and training runtime for execution review, replay drills, and robot-assisted trading experiments on separate infrastructure.")
    add_bullet(doc, "TWST Lab: Designed guardrails for environment isolation, persistent state, runtime checks, and Discord-driven operational workflows.")
    add_bullet(doc, "UltraTrading: Built a multi-runner crypto trading and paper-trading system with closed-loop learning, RPC safety wrappers, deploy verification, and audit layers.")
    add_bullet(doc, "UltraTrading: Focused on engineering reliability problems such as data drift, silent failures, schema-code alignment QA, reconciliation, and production-safe rollout.")

    add_section_heading(doc, "Professional Experience")
    add_role(
        doc,
        "Account Manager",
        "Cloudio Interactive Tech.",
        "2025 - 2026",
        [
            "Managed software project delivery, client requirements alignment, and coordination across internal and external development teams.",
            "Drove execution planning, issue tracking, and AI-assisted collaboration across product and engineering workflows.",
        ],
    )
    add_role(
        doc,
        "Technical Manager",
        "Softwidom Tech.",
        "2022 - 2024",
        [
            "Led cross-functional delivery for ERP/WMS, ESG carbon emission management, and AI prediction systems across frontend, backend, DBA, and PM functions.",
            "Translated business requirements into system architecture, delivery plans, and execution priorities.",
        ],
    )
    add_role(
        doc,
        "Technical Director",
        "BitSense IT Co., Ltd.",
        "2019 - 2022",
        [
            "Led product and engineering delivery for crypto wallet apps, NFT systems, Web3 staking products, and smart-contract-based platforms.",
            "Managed multidisciplinary teams across app, frontend, full-stack, UI/UX, and partner integrations.",
        ],
    )
    add_role(
        doc,
        "Product / Business Unit Leadership",
        "GoodWay Technology / Taiwan Anjiew / CastleNet / LeadTek / SinoPac Software",
        "2001 - 2018",
        [
            "20+ years of product and project leadership across electronics, networking, IoT, OEM/ODM, surveillance, and software systems.",
            "Delivered products from concept to shipment, including HTC VR-related ODM programs, smart home platforms, and enterprise software systems.",
        ],
    )

    add_section_heading(doc, "Education")
    add_bullet(doc, "Master's Degree, Information Management (MBA), Yuan Ze University, Taiwan | 2011 - 2012")
    add_bullet(doc, "Bachelor's Degree, Electrical Engineering, Guangwu Technical College, Taiwan")

    add_section_heading(doc, "Selected Achievements")
    add_bullet(doc, "7 patents across Taiwan, Japan, and Germany")
    add_bullet(doc, "Technical leadership across blockchain wallet, NFT, fintech, and AI-assisted software systems")
    add_bullet(doc, "Cross-domain background spanning hardware, software, and AI-native product delivery")

    return doc


if __name__ == "__main__":
    doc = build_cv()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")
